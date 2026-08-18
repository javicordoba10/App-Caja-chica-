import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def create_document():
    doc = Document()

    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Palette
    ORANGE = RGBColor(186, 72, 23)     # #BA4817
    DARK_BLUE = RGBColor(26, 36, 43)   # #1A242B
    GREY_TEXT = RGBColor(90, 95, 100)  # #5A5F64
    BLACK = RGBColor(0, 0, 0)

    # Base Normal Style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = DARK_BLUE

    # Document Header Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_title.add_run("PROPUESTA COMERCIAL Y VALORIZACIÓN DE SOFTWARE\n")
    run_sub.font.size = Pt(12)
    run_sub.font.bold = True
    run_sub.font.color.rgb = ORANGE

    run_main = p_title.add_run("SISTEMA DE GESTIÓN Y CONTROL DE CAJA CHICA EMPRESARIAL")
    run_main.font.size = Pt(20)
    run_main.font.bold = True
    run_main.font.color.rgb = DARK_BLUE

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_meta = p_meta.add_run("Multi-Plataforma (Web & Mobile Android/iOS) • Marca Blanca • Integración Firebase Cloud")
    r_meta.font.size = Pt(10)
    r_meta.font.italic = True
    r_meta.font.color.rgb = GREY_TEXT

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # 1. RESUMEN EJECUTIVO
    h1 = doc.add_heading(level=1)
    r = h1.add_run("1. Resumen Ejecutivo y Valor Tecnológico")
    r.font.color.rgb = ORANGE
    r.font.size = Pt(14)
    r.font.bold = True

    p = doc.add_paragraph(
        "El presente documento establece la propuesta comercial y esquemas de valorización para la plataforma de "
        "Gestión de Caja Chica y Egresos Empresariales. La solución ha sido concebida bajo estándares corporativos "
        "de alto rendimiento, utilizando un stack moderno (Flutter + Firebase Cloud Infrastructure) que garantiza "
        "multiplataforma nativa (Android, iOS y Web) con arquitectura de Marca Blanca (Multi-Tenant)."
    )
    p.paragraph_format.space_after = Pt(8)

    p_features = doc.add_paragraph()
    p_features.paragraph_format.space_after = Pt(12)
    runs = [
        ("Características Clave del Software:\n", True),
        ("• Captura y Auditoría de Comprobantes: ", True), ("Subida directa por cámara o galería con pre-visualización de PDF/imágenes e historial digitalizado.\n", False),
        ("• Seguridad & Autenticación de Dominio: ", True), ("Verificación estricta de correo electrónico vía Firebase Auth, control de usuarios activos y roles (Admin / Operador).\n", False),
        ("• Reglas de Optimización de Costos Nube: ", True), ("Regla de ciclo de vida automatizada (auto-eliminación de comprobantes adjuntos a los 30 días en Google Cloud Storage) preservando el registro contable histórico en Firestore sin costo extra.\n", False),
        ("• Multi-Empresa & Marca Blanca: ", True), ("Identidad visual dinámica (colores, logotipos y nombres configurables por cliente/unidad de negocio).\n", False),
        ("• Reportes y Exportación: ", True), ("Exportación rápida de balances, filtros por establecimiento y métodos de pago.", False)
    ]
    for text, is_bold in runs:
        r = p_features.add_run(text)
        if is_bold:
            r.font.bold = True
            r.font.color.rgb = DARK_BLUE
        else:
            r.font.color.rgb = GREY_TEXT

    # 2. ESQUEMAS DE MODELO DE NEGOCIO Y PRESUPUESTOS
    h2 = doc.add_heading(level=1)
    r = h2.add_run("2. Alternativas de Comercialización y Presupuestos")
    r.font.color.rgb = ORANGE
    r.font.size = Pt(14)
    r.font.bold = True

    doc.add_paragraph(
        "Se presentan tres (3) modelos de negocios estructurados para adaptarse a la estrategia comercial del cliente o inversor:"
    )

    # TABLE OF COMPARISON
    table = doc.add_table(rows=4, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    headers = ["Modelo de Venta", "Inversión Inicial", "Costo Recurrente", "Perfil de Cliente Ideal"]
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "BA4817")
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(10)

    rows_data = [
        ("Opción 1: Licencia Llave en Mano\n(Venta de App Completa)", "USD $3.500 - $5.000\n(Pago Único)", "Sin costo fijo\n(Servidor en su cuenta)", "Empresas corporativas que desean propiedad total del software y del código fuente."),
        ("Opción 2: Membresía SaaS\n(Suscripción Mensual)", "USD $0\n(Setup Bonificado)", "USD $35 - $60 / mes\n(por empresa o sucursal)", "PyMEs y empresas agrícolas que buscan baja inversión inicial y flexibilidad."),
        ("Opción 3: Modelo Híbrido\n(Licencia Base + Soporte)", "USD $1.200 - $1.800\n(Setup e Instalación)", "USD $20 - $30 / mes\n(Mantenimiento y Nube)", "Empresas que quieren su propia instancia configurada pero con soporte continuo.")
    ]

    for row_idx, data in enumerate(rows_data, start=1):
        row_cells = table.rows[row_idx].cells
        bg_color = "F8F9FA" if row_idx % 2 == 1 else "FFFFFF"
        for col_idx, text in enumerate(data):
            row_cells[col_idx].text = text
            set_cell_background(row_cells[col_idx], bg_color)
            set_cell_margins(row_cells[col_idx], top=100, bottom=100, left=120, right=120)
            p = row_cells[col_idx].paragraphs[0]
            if col_idx in [0, 1, 2]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.size = Pt(9.5)
                if col_idx == 0:
                    run.font.bold = True
                    run.font.color.rgb = DARK_BLUE

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # 3. DETALLE PROFUNDO DE LAS OPCIONES
    h3 = doc.add_heading(level=1)
    r = h3.add_run("3. Desglose Detallado de los 3 Modelos de Presupuesto")
    r.font.color.rgb = ORANGE
    r.font.size = Pt(14)
    r.font.bold = True

    # OPTION 1
    doc.add_heading(level=2, text="OPCIÓN 1: Venta de Licencia Perpetua / Código Fuente (Llave en Mano)")
    p1 = doc.add_paragraph()
    p1.add_run("• Descripción: ").bold = True
    p1.add_run("Transferencia completa de los derechos de uso o código fuente del software. La aplicación se despliega directamente en las cuentas de Firebase y servidores del comprador.\n")
    p1.add_run("• Inversión Sugerida: ").bold = True
    p1.add_run("USD $4.200 (Rango de mercado: USD $3.500 a USD $5.000).\n")
    p1.add_run("• Qué Incluye:\n").bold = True
    p1.add_run("   - Entrega del código fuente completo en Flutter (Web, Android, iOS).\n")
    p1.add_run("   - Despliegue e instalación en el proyecto de Firebase del cliente.\n")
    p1.add_run("   - Compilación del ejecutable APK Release para dispositivos Android.\n")
    p1.add_run("   - 30 días de soporte y garantía pos-entrega para correcciones.\n")
    p1.add_run("• Infraestructura Nube: ").bold = True
    p1.add_run("A cargo del cliente (en Plan Gratuito Spark de Firebase los costos suelen ser $0 USD para operaciones estándar).")

    # OPTION 2
    doc.add_heading(level=2, text="OPCIÓN 2: Modelo SaaS (Membresía Mensual / Recurrente)")
    p2 = doc.add_paragraph()
    p2.add_run("• Descripción: ").bold = True
    p2.add_run("Alquiler del servicio en la nube (Software as a Service). El cliente paga un canon mensual por usar la plataforma alojada y administrada por usted.\n")
    p2.add_run("• Estructura de Precios Sugerida:\n").bold = True
    p2.add_run("   - Plan PyME (hasta 5 usuarios / 1 sucursal): ").bold = True
    p2.add_run("USD $35 / mes (aprox. $40.000 - $45.000 ARS/mes).\n")
    p2.add_run("   - Plan Corporativo (hasta 20 usuarios / múltiples sucursales): ").bold = True
    p2.add_run("USD $75 / mes.\n")
    p2.add_run("• Qué Incluye:\n").bold = True
    p2.add_run("   - Acceso 24/7 a la versión Web y aplicación Android.\n")
    p2.add_run("   - Infraestructura y servidores Nube (Firebase) 100% incluidos.\n")
    p2.add_run("   - Actualizaciones continuas, copias de seguridad e infraestructura garantizada.\n")
    p2.add_run("   - Soporte técnico continuo vía email/WhatsApp.\n")
    p2.add_run("• Ventaja Comercial: ").bold = True
    p2.add_run("Genera ingresos recurrentes estables (ARR/MRR) a largo plazo.")

    # OPTION 3
    doc.add_heading(level=2, text="OPCIÓN 3: Modelo Híbrido (Implementación Inicial + Abono de Mantenimiento)")
    p3 = doc.add_paragraph()
    p3.add_run("• Descripción: ").bold = True
    p3.add_run("Combinación equilibrada. Se cobra un costo inicial de configuración e integración personalizada (Setup), más una cuota mensual reducida que cubre soporte y servidores Nube.\n")
    p3.add_run("• Estructura de Precios Sugerida:\n").bold = True
    p3.add_run("   - Setup Inicial de Personalización e Instalación: ").bold = True
    p3.add_run("USD $1.500 (pago único).\n")
    p3.add_run("   - Abono Mensual de Mantenimiento & Nube: ").bold = True
    p3.add_run("USD $25 / mes.\n")
    p3.add_run("• Qué Incluye:\n").bold = True
    p3.add_run("   - Configuración inicial de logo, colores de marca e importación de usuarios.\n")
    p3.add_run("   - Mantenimiento técnico preventivo y correctivo.\n")
    p3.add_run("   - Alojamiento de datos y regla de auto-limpieza de fotos en Google Cloud Storage.")

    # 4. RECOMENDACIÓN ESTRATÉGICA
    h4 = doc.add_heading(level=1)
    r = h4.add_run("4. Recomendación del Especialista en Ventas")
    r.font.color.rgb = ORANGE
    r.font.size = Pt(14)
    r.font.bold = True

    p_rec = doc.add_paragraph()
    p_rec.add_run("Si el objetivo es construir un negocio escalable con valor financiero creciente, el ").font.size = Pt(11)
    r_bold = p_rec.add_run("Modelo SaaS (Opción 2) o Híbrido (Opción 3)")
    r_bold.font.bold = True
    r_bold.font.color.rgb = ORANGE
    p_rec.add_run(
        " son los más recomendados. Un cliente que paga USD $35/mes genera USD $420/año. "
        "Con 10 empresas clientes, se obtienen USD $4.200 anuales de forma recurrente, superando con creces la venta única de la Opción 1 "
        "y reteniendo la propiedad intelectual de la plataforma."
    )

    doc.save("Propuesta_Comercial_Caja_Chica.docx")
    print("Documento guardado con éxito: Propuesta_Comercial_Caja_Chica.docx")

if __name__ == '__main__':
    create_document()
